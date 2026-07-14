"""
Document parser — handles PDF, DOCX, Markdown, and plain text.
Returns a list of RawDocument objects ready for chunking.
"""

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class RawDocument:
    filename: str
    content: str                    # Full extracted text
    file_type: str                  # pdf | docx | md | txt
    page_count: Optional[int] = None


class DocumentParser:
    """
    Parses uploaded files into RawDocument objects.
    Supports: PDF (PyMuPDF), DOCX (python-docx), MD/TXT (native).
    """

    SUPPORTED_TYPES = {".pdf", ".docx", ".md", ".txt"}

    def parse(self, filename: str, file_bytes: bytes) -> RawDocument:
        suffix = Path(filename).suffix.lower()

        if suffix not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: {self.SUPPORTED_TYPES}")

        logger.info("parsing_document", filename=filename, type=suffix)

        if suffix == ".pdf":
            return self._parse_pdf(filename, file_bytes)
        elif suffix == ".docx":
            return self._parse_docx(filename, file_bytes)
        else:
            return self._parse_text(filename, file_bytes, suffix.lstrip("."))

    def _parse_pdf(self, filename: str, file_bytes: bytes) -> RawDocument:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = []
            for page in doc:
                pages.append(page.get_text())
            content = "\n\n".join(pages)
            return RawDocument(
                filename=filename,
                content=content,
                file_type="pdf",
                page_count=len(doc)
            )
        except ImportError:
            raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")

    def _parse_docx(self, filename: str, file_bytes: bytes) -> RawDocument:
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            return RawDocument(
                filename=filename,
                content="\n\n".join(paragraphs),
                file_type="docx"
            )
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    def _parse_text(self, filename: str, file_bytes: bytes, file_type: str) -> RawDocument:
        content = file_bytes.decode("utf-8", errors="replace")
        return RawDocument(
            filename=filename,
            content=content,
            file_type=file_type
        )
