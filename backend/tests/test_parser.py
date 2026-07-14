import io

import pytest
from docx import Document as DocxDocument

from ingestion.parser import DocumentParser, RawDocument


@pytest.fixture
def parser():
    return DocumentParser()


def test_parse_txt(parser):
    doc = parser.parse("notes.txt", b"hello world")
    assert isinstance(doc, RawDocument)
    assert doc.content == "hello world"
    assert doc.file_type == "txt"
    assert doc.filename == "notes.txt"


def test_parse_md(parser):
    doc = parser.parse("spec.md", b"# Title\n\nBody text")
    assert doc.content == "# Title\n\nBody text"
    assert doc.file_type == "md"


def test_parse_text_decodes_invalid_utf8_with_replacement(parser):
    doc = parser.parse("bad.txt", b"valid \xff\xfe bytes")
    assert "�" in doc.content  # replacement character, no exception raised


def test_parse_unsupported_extension_raises(parser):
    with pytest.raises(ValueError, match="Unsupported file type"):
        parser.parse("archive.zip", b"whatever")


def test_parse_docx(parser):
    docx_doc = DocxDocument()
    docx_doc.add_paragraph("First paragraph")
    docx_doc.add_paragraph("Second paragraph")
    table = docx_doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"

    buf = io.BytesIO()
    docx_doc.save(buf)

    parsed = parser.parse("plan.docx", buf.getvalue())
    assert parsed.file_type == "docx"
    assert "First paragraph" in parsed.content
    assert "Second paragraph" in parsed.content
    assert "A | B" in parsed.content


def test_parse_pdf(parser):
    fitz = pytest.importorskip("fitz")
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((72, 72), "Hello from PDF")
    pdf_bytes = pdf_doc.tobytes()
    pdf_doc.close()

    parsed = parser.parse("report.pdf", pdf_bytes)
    assert parsed.file_type == "pdf"
    assert parsed.page_count == 1
    assert "Hello from PDF" in parsed.content
