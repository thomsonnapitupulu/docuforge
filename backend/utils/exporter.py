"""
Exporter — converts the compiled markdown document into a .docx file.

Handles:
- H1 → Document Title style
- H2 → Heading 1 style
- H3 → Heading 2 style
- Bullet lists → List Bullet style
- Code blocks → preformatted Courier New
- Tables → Word table with header row shading
- Bold/italic inline text
"""

import io
import re
from docx import Document
from docx.shared import Pt
import structlog

logger = structlog.get_logger()


def markdown_to_docx(markdown_text: str, title: str = "Document") -> bytes:
    """
    Convert a markdown string into a .docx binary payload.
    Returns the raw bytes of the Word document.
    """
    doc = Document()

    # ── Document styling ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)

        # ── Code blocks ───────────────────────────────────────────────────────
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph(code_text, style="No Spacing")
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9)

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60, style="No Spacing")

        # ── Bullet lists ──────────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            _add_inline_paragraph(doc, line[2:].strip(), style_name="List Bullet")

        # ── Numbered lists ────────────────────────────────────────────────────
        elif re.match(r"^\d+\. ", line):
            content = re.sub(r"^\d+\. ", "", line).strip()
            _add_inline_paragraph(doc, content, style_name="List Number")

        # ── Tables ────────────────────────────────────────────────────────────
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # ── Blockquotes / callouts ────────────────────────────────────────────
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip(), style="Quote")

        # ── Regular paragraph ─────────────────────────────────────────────────
        elif line.strip():
            _add_inline_paragraph(doc, line.strip())

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("docx_export_complete", title=title)
    return buf.read()


def _add_inline_paragraph(doc: Document, text: str, style_name: str = "Normal"):
    """Add a paragraph with inline bold/italic markdown rendering."""
    p = doc.add_paragraph(style=style_name)
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            p.add_run(part)


def _add_table(doc: Document, table_lines: list[str]):
    """Convert markdown table lines into a Word table."""
    rows = []
    for line in table_lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue  # Skip separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < max_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text
                if row_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
